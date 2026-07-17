from __future__ import annotations

import argparse
import csv
from copy import deepcopy
import hashlib
from pathlib import Path
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips


SOURCE_SHA256 = "886e511c233f4002a039ba533df86f697d81703585f3d2bc032bb6b5f812560e"
FIGURE1_MEDIA_SHA256 = "6d65e98734cb94851b545f11d744c32535fd00efb6318a1de203aa28333e5e84"
EXPECTED_FIGURE_SHA256 = {
    "Fig2_camera_ready_contamination.png": "28911248c57dec3ce1815b1d6ca8820d0ad5ab73ba40be2a5edbe3d32e9b3320",
    "Fig3_camera_ready_audit.png": "15c12949098c863a462818367390121ed78b9a21cb878c567cea190afdac2e8a",
    "Fig4_camera_ready_cdf.png": "111ec47f0d060ffd4ac80d343ffd443ec8c7f7a61a639c0ddbc4b7f36c085a2d",
    "Fig5_camera_ready_bo_lhs.png": "f517d266ef53348662411e9682c59582caae14458becc99ffa099a1ae42b3185",
    "Table_I_camera_ready_ablation.png": "8d81cc4404d93ea2da988223cb71d47a90168832eb2e03099da718bc27b4dc34",
}


ABSTRACT = (
    "Urban bus-corridor digital twins often ingest door-to-door observations that mix traffic motion "
    "with holding, layover, and stop service. We present RCMDT, an operator-aware two-level calibration "
    "protocol that separates a bus/stop parameter vector from a background-traffic context and freezes "
    "each layer before cross-day evaluation. A fixed audit flags link-hour medians satisfying T > 325 s, "
    "effective speed < 5 km/h, and distance ≤ 1500 m; development-only sensitivity, MAD, and Isolation "
    "Forest provide robustness checks without using IRN to tune thresholds. L1 compares Gaussian-process "
    "expected-improvement search with continued Latin-hypercube sampling under a common 40-evaluation "
    "budget; L2 uses 10 ensemble members and three iterations with moving-only observations. On Dec. 19 "
    "development and Dec. 30 cross-day data, Rule C retains 56.6% and 47.6% of eligible link keys and yields "
    "lower full-window speed K-S distances than the two fitted alternatives, at the cost of lower retention. "
    "Across four common ablation seeds, Full-RCMDT obtains cross-day K-S 0.348 ± 0.035, close to Raw-RCMDT "
    "at 0.351 ± 0.027; BO-only is lowest at 0.339 ± 0.035. BO has a lower mean final L1 score than LHS and "
    "wins four of five seeds, but no configuration is uniformly best. The results support an auditable freeze "
    "protocol and expose remaining transfer and worst-window discrepancies."
)


TEXT = {
    5: (
        "Bus-corridor observations combine vehicle motion, passenger service, schedule recovery, holding, and "
        "terminal layover. When these mechanisms are collapsed into a single door-to-door (D2D) travel time, a "
        "calibrator can attribute operational delay to traffic or vehicle parameters. The resulting digital twin "
        "may match a mean statistic while misrepresenting tails and cross-day behavior. Digital-twin research "
        "therefore requires both an executable model and an explicit observation operator [1], [2]."
    ),
    6: (
        "GPS-derived distances and speeds introduce further bias [3], while traffic-state estimation and iterative "
        "ensemble methods assume that observations and model outputs describe compatible quantities [4]-[6]. Bus "
        "control studies show that holding and regulation reshape passenger-experienced time [7]-[9]. Existing work "
        "does not establish how a bus-corridor twin should separate those semantics, freeze vehicle scopes, and test "
        "the result under a different calendar day."
    ),
    7: (
        "RCMDT addresses this gap with three contracts. First, an auditable rule separates long-duration, "
        "near-stationary, short-link records before L2 assimilation. Second, L1 calibrates bus and stop-service "
        "parameters, whereas L2 updates only background-traffic context; each layer is frozen outside its assigned "
        "stage. Third, validation uses common link keys, event-level K-S distance, worst 15-min windows, and a "
        "Dec. 30 cross-day split without re-optimization."
    ),
    8: (
        "The contribution is the auditable protocol rather than a new BO or smoother algorithm. Its evidence chain "
        "tests contamination, threshold stability, statistical and fitted audit alternatives, an A0-A4 mechanism "
        "ablation, and equal-budget BO versus LHS. IRN supplies an external plausibility check only [12], [13]; it "
        "does not set thresholds or act as semantic ground truth."
    ),
    13: "Relevant work falls into five streams.",
    15: (
        "Digital-twin surveys emphasize synchronization, uncertainty, and a maintained data-to-model update path "
        "[1], [2]. RCMDT narrows this general requirement to a reproducible bus-corridor calibration and freeze "
        "protocol."
    ),
    17: (
        "GPS reconstruction can inflate distance and distort speed and travel-time tails [3]. For buses, the same "
        "D2D interval can also include stop service and operator intervention, so the observation definition must be "
        "logged with the calibration target."
    ),
    19: (
        "Ensemble filtering and iterative smoothing recover traffic states from mobile observations [4]-[6]. Their "
        "validity depends on semantic compatibility. RCMDT therefore limits L2 to moving-only observations and "
        "freezes bus/stop parameters during background-context updates."
    ),
    21: (
        "Holding, scheduling, and control policies alter observed bus travel times [7]-[9]. Distributionally robust "
        "optimization motivates shift-sensitive evaluation [10], but a low K-S statistic on one split alone does not "
        "establish generalization or semantic truth."
    ),
    24: (
        "We represent the simulator as the stochastic mapping in (1), driven by a bus/stop vector, a "
        "background-traffic context, and simulation randomness."
    ),
    28: (
        "For each valid candidate, cumulative observed and simulated arrival times are matched at downstream stops. "
        "The RMSE component is defined in (2)."
    ),
    30: (
        "For joined downstream stop i, e_i is the mean simulated cumulative arrival time minus the mean observed "
        "cumulative arrival time, both relative to the first matched stop. Observations are aggregated by route, "
        "direction, and link sequence; the origin is excluded and n must be at least three. One deterministic SUMO "
        "seed is derived from the optimization seed and evaluation index."
    ),
    32: (
        "Validation reports the two-sample K-S statistic D in (3) between event-level observed and simulated samples "
        "on the same frozen link-key set. It is a distance, not a p-value or pass/fail label."
    ),
    34: (
        "Rule C flags an eligible link-hour median only when T > 325 s, effective speed < 5 km/h, and distance ≤ "
        "1500 m. Thresholds are predeclared; Dec. 19 sensitivity cannot retune them after Dec. 30 is inspected. "
        "Worst-window K-S is the maximum valid 15-min statistic on a 60-s stride."
    ),
    37: (
        "Fig. 1 shows the accepted two-level flow. L1 estimates bus/stop parameters, L2 reconciles background "
        "context, and the logbook records the observation audit and freeze state. Calibration and validation remain "
        "separate, preventing a validation discrepancy from silently changing thresholds or parameters."
    ),
    38: (
        "This scope separation differentiates RCMDT from an undifferentiated dual loop: bus car-following and stop "
        "service remain fixed during L2, background traffic remains fixed during L1, and raw versus moving-only L2 "
        "semantics are explicit ablation mechanisms. Each ablation row therefore maps to a declared mechanism; the "
        "executable definitions in Sections III-B-C govern the schematic's shorthand labels."
    ),
    44: (
        "The GP is trained on pairs of evaluated bus/stop candidates and their scalar L1 scores. A new candidate z "
        "is one complete in-bounds bus/stop parameter vector. The surrogate supplies predictive mean μ(z) and "
        "standard deviation σ(z); the next unevaluated vector maximizes expected improvement (4), with "
        "standardized improvement (5) [14], [15]."
    ),
    46: (
        "Here f* is the best feasible observed score, Φ and φ are the standard normal CDF and density, and "
        "γ = (f* - μ(z))/σ(z). Candidate selection is therefore arg max EI over the generated feasible "
        "candidate pool, followed by one deterministic simulation."
    ),
    48: (
        "The 68X objective (6) combines RMSE, MAE, dispersion of absolute errors, and the 0.9 quantile. Route 960 is "
        "a feasibility anchor rather than a pooled error term."
    ),
    50: (
        "Weights are α = 1, λ = 0.5, and β = 0.3. A valid candidate is feasible when 960 RMSE ≤ 350 s; "
        "a constraint violation receives the declared penalty, while missing output triggers deterministic retry and "
        "is never inserted as a numerical observation. BO and continued LHS each use 40 successful evaluations."
    ),
    52: (
        "L2 applies the iterative analysis update (7) while the bus/stop vector is frozen. Ensemble covariances "
        "approximate the Jacobian, and bounds are enforced after each damped update."
    ),
    55: "The released protocol uses Nₑ = 10, K = 3 iterations, and initial damping 0.3.",
    56: (
        "L2 receives moving-only link speeds. Supplying raw D2D speed embeds holding and stop service in a "
        "background-flow target; A3 preserves that mismatch deliberately, while A2 and A4 use the audited semantic."
    ),
    58: (
        "MAD and Isolation Forest are fitted jointly on eligible Dec. 19 link-hour records and frozen before Dec. 30. "
        "All methods begin with the same eligible keys and fixed A0 simulation. Retention, full-window K-S, worst-window "
        "K-S, and IRN contradiction counts are reported; the latter is diagnostic because matched denominators are small."
    ),
    61: (
        "The SUMO testbed covers KMB routes 68X and 960. Development uses Dec. 19, 2025, 17:00-18:00 HKT; cross-day "
        "evaluation uses Dec. 30, 2025, 15:00-16:00 HKT. No held-out same-day test exists."
    ),
    62: (
        "D2D observations provide travel-time and effective-speed events. Processed IRN segment speeds are matched "
        "only for external plausibility checks [12], [13] and never enter calibration, threshold fitting, or scoring."
    ),
    64: (
        "A0 is zero-shot; A1 enables L1 only; A2 enables L2 only; A3 enables both layers with raw D2D supplied to L2; "
        "A4 enables both layers with moving-only L2 input. Four common successful seeds are reported. Every row uses "
        "the same Rule-C-clean evaluation population; full-window K-S requires at least 20 events per source and each "
        "15-min subwindow requires five."
    ),
    65: (
        "L1 uses 15 shared LHS evaluations followed by 25 BO or continued-LHS evaluations. L2 starts from common "
        "priors and ensemble seeds. Effective input hashes verify the intended mechanism equalities and differences."
    ),
    68: (
        "Fig. 2 documents the contamination chain: 74 of 200 development events are flagged, leaving 126 clean "
        "events; the link-hour geometry and cumulative trajectories show why low D2D speed need not imply traffic "
        "congestion. Fig. 3 evaluates the rule rather than treating it as ground truth."
    ),
    69: (
        "Rule C retains 0.566 of eligible development keys and 0.476 cross-day, with full-window K-S 0.175 and 0.368. "
        "MAD retains all keys and yields 0.417/0.454; Isolation Forest retains 0.961/0.794 and yields 0.417/0.444. "
        "Rule C is lower on full-window K-S but not uniformly on worst-window K-S; IRN denominators remain too small "
        "for an accuracy claim."
    ),
    75: (
        "Table I reports mean ± sample SD over four common seeds. On Dec. 19, A0 has the lowest K-S (0.154 ± "
        "0.003). On Dec. 30, A1 is lowest (0.339 ± 0.035), followed by A4 (0.348 ± 0.035) and A3 (0.351 ± "
        "0.027); fixed real sample counts are 60 and 29."
    ),
    76: (
        "A4 improves cross-day K-S relative to A0 (0.390 ± 0.009) but does not dominate development or worst-window "
        "metrics; its cross-day worst-window value is 0.875 ± 0.050. The ablation supports mechanism separation and "
        "also exposes residual short-window sensitivity."
    ),
    81: (
        "Fig. 4 shows the Rule-C-clean speed CDF against each A4 seed on both dates. The spread across simulation "
        "curves and the Dec. 30 tail show remaining structural discrepancy; no equality or universal transfer claim "
        "is made."
    ),
    85: (
        "Fig. 5 compares equal budgets across five optimization seeds. Final feasible L1 scores average 1988 ± 522 "
        "for BO and 2095 ± 387 for LHS; BO is lower in four seeds and reaches the predeclared target in three seeds "
        "versus two for LHS. Seed 3 favors LHS, so the evidence supports efficiency on average, not dominance."
    ),
    89: (
        "Rule C trades retention for a lower full-window discrepancy, and its IRN check has sparse matched flagged "
        "records. Only four common ablation seeds completed; A4 is not best on every metric, and worst-window K-S "
        "remains high. The context vector contains calibrated background parameters rather than directly observed "
        "traffic states. These limits constrain the result to an auditable cross-day protocol on two routes."
    ),
    91: (
        "RCMDT separates bus/stop calibration, background-context reconciliation, and observation auditing under a "
        "recorded freeze protocol. The revised evidence shows that Rule C removes a physically interpretable "
        "long-duration, near-stationary, short-link regime and produces lower full-window K-S than MAD and Isolation "
        "Forest on both dates, while retaining fewer keys. The A0-A4 study does not identify a uniformly superior "
        "configuration: Full-RCMDT improves cross-day K-S over zero-shot and is close to Raw-RCMDT, whereas BO-only "
        "is lowest and worst-window discrepancies persist. Under equal 40-evaluation budgets, BO lowers the mean "
        "final feasible L1 score and wins four of five seeds, with one clear LHS counterexample. The contribution is "
        "therefore a reproducible semantic and scope contract that reveals both gains and failure modes. Future work "
        "should enlarge the cross-day sample, improve IRN coverage, and model holding and schedule control explicitly."
    ),
    93: (
        "The authors acknowledge the use of Google Gemini 3 Pro and OpenAI GPT-5.2 Codex in code development, "
        "experiment scripting, and language editing. The authors verified the reported data, analyses, and final text."
    ),
}


HEADING_TEXT = {
    43: "B. L1: Gaussian-Process Search",
    51: "C. L2: Frozen Background-Context Update",
    57: "D. Observation Audit and Freeze Protocol",
    60: "A. Testbed and Data Sources",
    63: "B. Common Experimental Protocol",
    67: "A. Audit Stability and Alternatives",
    74: "B. Common-Protocol A0-A4 Ablation",
    80: "C. Cross-Day Distributional Evidence",
    84: "D. Equal-Budget Search",
    88: "E. Limitations and Interpretation",
}


FIGURE_CAPTIONS = [
    (
        "(a) Empirical speed CDFs for 200 raw and 126 Rule-C-clean development events. (b) Link-hour audit geometry: "
        "T > 325 s, effective speed < 5 km/h, and distance ≤ 1500 m; the annotation is placed outside the data cloud. "
        "(c) Observed traffic-only and simulated full/traffic-only cumulative trajectories."
    ),
    (
        "Audit robustness. (a)-(b) Dec. 19 sensitivity of retention and speed K-S; the outlined cell is the predeclared "
        "Rule C point. (c) Full-window and worst-15-min K-S for fixed, MAD, and Isolation-Forest audits. (d) Retention "
        "and IRN contradiction diagnostics with matched/unmatched counts; IRN is not a classification label."
    ),
    (
        "Speed CDFs for (a) Dec. 19 development and (b) Dec. 30 cross-day evaluation. Solid curves are the frozen "
        "Rule-C-clean observations (n = 60 and 29); dashed curves are A4 simulations for four common seeds."
    ),
    (
        "Equal-budget L1 search over five optimization seeds. Curves show the mean feasible cumulative-best score, "
        "bands show ± one sample SD, the vertical line ends the 15 shared LHS evaluations, and the horizontal line "
        "is the mean predeclared target. Lower is better."
    ),
]


REFERENCE_ADDITIONS = [
    (
        "J. Snoek, H. Larochelle, and R. P. Adams, \"Practical Bayesian Optimization of Machine Learning "
        "Algorithms,\" in Advances in Neural Information Processing Systems 25, 2012, pp. 2951–2959."
    ),
    "P. I. Frazier, \"A Tutorial on Bayesian Optimization,\" arXiv:1807.02811, 2018.",
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def clear_paragraph_content(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_text(paragraph, text: str) -> None:
    first_rpr = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        first_rpr = deepcopy(paragraph.runs[0]._r.rPr)
    clear_paragraph_content(paragraph)
    run = paragraph.add_run(text)
    if first_rpr is not None:
        run._r.insert(0, first_rpr)


def set_abstract(paragraph, text: str) -> None:
    clear_paragraph_content(paragraph)
    label = paragraph.add_run("Abstract—")
    label.bold = True
    label.italic = True
    paragraph.add_run(" " + text)


def add_math_symbol(paragraph, base: str, subscript: str | None = None):
    run = paragraph.add_run(base)
    run.italic = True
    if subscript is not None:
        sub = paragraph.add_run(subscript)
        sub.font.subscript = True
    return run


def set_definition_paragraph(paragraph) -> None:
    clear_paragraph_content(paragraph)
    paragraph.add_run("where ")
    add_math_symbol(paragraph, "θ", "bus")
    paragraph.add_run(" = (")
    add_math_symbol(paragraph, "t", "board")
    paragraph.add_run(", ")
    add_math_symbol(paragraph, "t", "fixed")
    paragraph.add_run(", τ, σ, ")
    add_math_symbol(paragraph, "minGap", "bus")
    paragraph.add_run(", accel, decel) contains stop-service and bus car-following parameters. ")
    add_math_symbol(paragraph, "x", "corr")
    paragraph.add_run(" = (capacityFactor, ")
    add_math_symbol(paragraph, "minGap", "background")
    paragraph.add_run(", impatience) contains only background-traffic context, and ξ denotes simulation "
                      "randomness. The bus/stop vector is frozen during L2 and validation.")


def set_background_context_paragraph(paragraph) -> None:
    clear_paragraph_content(paragraph)
    paragraph.add_run("Here ")
    add_math_symbol(paragraph, "x", "corr,a")
    paragraph.add_run(" and ")
    add_math_symbol(paragraph, "x", "corr,f")
    paragraph.add_run(" are the analyzed and forecast background-context vectors, respectively; ")
    add_math_symbol(paragraph, "y")
    paragraph.add_run(" contains moving-only link speeds, and ")
    add_math_symbol(paragraph, "h")
    paragraph.add_run(" maps context to those speeds. ")
    add_math_symbol(paragraph, "P", "f")
    paragraph.add_run(" and ")
    add_math_symbol(paragraph, "R")
    paragraph.add_run(" are the ensemble forecast and observation-error covariances. Only capacityFactor, ")
    add_math_symbol(paragraph, "minGap", "background")
    paragraph.add_run(", and impatience are updated; ")
    add_math_symbol(paragraph, "minGap", "bus")
    paragraph.add_run(" and every component of ")
    add_math_symbol(paragraph, "θ", "bus")
    paragraph.add_run(" remain frozen.")


def clone_with_text(paragraph, text: str):
    element = deepcopy(paragraph._p)
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)
    run = OxmlElement("w:r")
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run.append(deepcopy(paragraph.runs[0]._r.rPr))
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    element.append(run)
    return element


def set_page_break_before(element) -> None:
    ppr = element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        element.insert(0, ppr)
    if ppr.find(qn("w:pageBreakBefore")) is None:
        ppr.append(OxmlElement("w:pageBreakBefore"))


def set_section_break_continuous(element) -> None:
    sect_pr = element.find("./w:pPr/w:sectPr", element.nsmap)
    if sect_pr is None:
        raise RuntimeError("Expected a section-break paragraph")
    section_type = sect_pr.find(qn("w:type"))
    if section_type is None:
        section_type = OxmlElement("w:type")
        sect_pr.insert(0, section_type)
    section_type.set(qn("w:val"), "continuous")


def set_keep_with_next(element) -> None:
    ppr = element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        element.insert(0, ppr)
    if ppr.find(qn("w:keepNext")) is None:
        ppr.append(OxmlElement("w:keepNext"))


def picture_element(document: Document, path: Path, width_inches: float):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    paragraph.add_run().add_picture(str(path), width=Inches(width_inches))
    element = paragraph._p
    element.getparent().remove(element)
    set_keep_with_next(element)
    return element


def set_cell_width(cell, width_twips: int) -> None:
    cell.width = Twips(width_twips)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def native_table_element(document: Document, csv_path: Path):
    with csv_path.open("r", encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle))
    if len(rows) != 5:
        raise RuntimeError(f"Expected five Table I rows, found {len(rows)}")

    headers = (
        "Configuration",
        "Dev. K-S",
        "Dev. worst",
        "Cross-day K-S",
        "Cross-day worst",
        "Sim. n (dev.)",
        "Sim. n (cross-day)",
    )
    widths = (1600, 1360, 1360, 1440, 1440, 1440, 1440)
    total_width = sum(widths)

    def mean_sd(row, mean_key: str, std_key: str) -> str:
        return f"{float(row[mean_key]):.3f} ± {float(row[std_key]):.3f}"

    values = []
    for row in rows:
        values.append(
            (
                f"{row['config_id']} {row['configuration']}",
                mean_sd(row, "ks_speed_development_mean", "ks_speed_development_std"),
                mean_sd(row, "worst_15min_ks_development_mean", "worst_15min_ks_development_std"),
                mean_sd(row, "ks_speed_cross_day_mean", "ks_speed_cross_day_std"),
                mean_sd(row, "worst_15min_ks_cross_day_mean", "worst_15min_ks_cross_day_std"),
                mean_sd(row, "n_sim_development_mean", "n_sim_development_std"),
                mean_sd(row, "n_sim_cross_day_mean", "n_sim_cross_day_std"),
            )
        )

    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "666666")
        borders.append(border)
    tbl_pr.append(borders)

    cell_margins = OxmlElement("w:tblCellMar")
    for side, value in (("top", 35), ("left", 45), ("bottom", 35), ("right", 45)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        cell_margins.append(node)
    tbl_pr.append(cell_margins)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row_values in enumerate((headers, *values)):
        row = table.rows[0] if row_index == 0 else table.add_row()
        if row_index == 0:
            tr_pr = row._tr.get_or_add_trPr()
            tr_pr.append(OxmlElement("w:tblHeader"))
        for column_index, (cell, text) in enumerate(zip(row.cells, row_values, strict=True)):
            set_cell_width(cell, widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shade = OxmlElement("w:shd")
                shade.set(qn("w:fill"), "D9D9D9")
                cell._tc.get_or_add_tcPr().append(shade)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
            paragraph.paragraph_format.line_spacing = 1
            run = paragraph.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(7)
            run.bold = row_index == 0

    element = table._tbl
    element.getparent().remove(element)
    return element


def table_note_element(document: Document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(
        "Mean ± sample SD over four common seeds. Fixed real events: development n=60; cross-day n=29."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(6.5)
    element = paragraph._p
    element.getparent().remove(element)
    return element


def replace_table_title_suffix(paragraph) -> None:
    element = paragraph._p
    field_end_index = None
    for index, child in enumerate(element):
        field_char = child.find(qn("w:fldChar"))
        if field_char is not None and field_char.get(qn("w:fldCharType")) == "end":
            field_end_index = index
            break
    if field_end_index is None:
        raise RuntimeError("Table I SEQ field end was not found")
    for child in list(element)[field_end_index + 1 :]:
        element.remove(child)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    small_caps = OxmlElement("w:smallCaps")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "16")
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "16")
    rpr.extend([small_caps, size, size_cs])
    run.append(rpr)
    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    text.text = ".  COMMON-PROTOCOL A0-A4 ABLATION"
    run.append(text)
    element.append(run)
    set_keep_with_next(element)


def insert_after(cursor, elements):
    for element in elements:
        cursor.addnext(element)
        cursor = element
    return cursor


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def update_equations(paragraphs) -> None:
    eq1 = paragraphs[25]._p.xpath(".//m:t|.//w:t")
    eq1[5].text = "bus"
    eq1[7].text = "x"
    eq7 = paragraphs[53]._p.xpath(".//m:t|.//w:t")
    eq7[1].text = "corr,a"
    eq7[4].text = "corr,f"
    eq7[11].text = "corr,f"


def add_update_fields_setting(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def embedded_hash(docx_path: Path, member: str) -> str:
    with zipfile.ZipFile(docx_path) as archive:
        return hashlib.sha256(archive.read(member)).hexdigest()


def build(source: Path, figures: Path, table_csv: Path, output: Path) -> None:
    if sha256(source) != SOURCE_SHA256:
        raise RuntimeError("Source manuscript hash differs from the approved baseline")
    for name, expected in EXPECTED_FIGURE_SHA256.items():
        actual = sha256(figures / name)
        if actual != expected:
            raise RuntimeError(f"Figure hash mismatch for {name}: {actual}")
    if not table_csv.is_file():
        raise RuntimeError(f"Table I source CSV does not exist: {table_csv}")

    document = Document(source)
    if len(document.paragraphs) != 108 or len(document.tables) != 1 or len(document.sections) != 3:
        raise RuntimeError("Unexpected source manuscript structure")
    p = list(document.paragraphs)

    set_abstract(p[3], ABSTRACT)
    for index, text in TEXT.items():
        set_text(p[index], text)
    for index, text in HEADING_TEXT.items():
        set_text(p[index], text)
    set_definition_paragraph(p[26])
    set_background_context_paragraph(p[54])
    p[31].style = document.styles["Heading 2"]
    set_text(p[31], "C. Distributional Evidence and Audit Contract")
    update_equations(p)

    # Remove redundant introduction prose after the concise four-paragraph version.
    for index in (9, 10, 11):
        remove_element(p[index]._p)

    # Add the missing fifth related-work stream using the manuscript's existing IEEE styles.
    related_heading = clone_with_text(p[20], "E. Simulation Calibration and Bayesian Optimization")
    related_body = clone_with_text(
        p[21],
        "Bayesian optimization uses a probabilistic surrogate to allocate expensive simulator calls [14], [15]. "
        "RCMDT does not claim a new acquisition rule; it contributes an equal-budget comparison, deterministic seed "
        "schedule, cross-route feasibility constraint, and a frozen handoff from bus/stop calibration to background "
        "reconciliation.",
    )
    p[22]._p.addprevious(related_heading)
    p[22]._p.addprevious(related_body)

    # Preserve Figure 1 and both section-break patterns exactly; only surrounding prose changes.

    # Save caption templates before removing the superseded visual block.
    caption_sources = [p[71], p[73], p[83], p[87]]
    new_captions = [
        clone_with_text(source_caption, caption_text)
        for source_caption, caption_text in zip(caption_sources, FIGURE_CAPTIONS, strict=True)
    ]

    # Remove the old Figure 2-Figure 5 paragraphs and the old editable table.
    for index in (70, 71, 72, 73, 78, 79, 82, 83, 86, 87):
        remove_element(p[index]._p)
    remove_element(document.tables[0]._tbl)

    # Retain the existing Table I SEQ field and IEEE title formatting, then move it into the gallery.
    table_title = p[77]._p
    remove_element(table_title)
    replace_table_title_suffix(p[77])

    fig2 = picture_element(document, figures / "Fig2_camera_ready_contamination.png", 7.0)
    fig3 = picture_element(document, figures / "Fig3_camera_ready_audit.png", 7.0)
    fig4 = picture_element(document, figures / "Fig4_camera_ready_cdf.png", 7.0)
    fig5 = picture_element(document, figures / "Fig5_camera_ready_bo_lhs.png", 3.5)
    table = native_table_element(document, table_csv)
    table_note = table_note_element(document)

    def wide_burst(after, content):
        start = deepcopy(p[39]._p)
        set_section_break_continuous(start)
        end = deepcopy(p[42]._p)
        return insert_after(after, [start, *content, end])

    # Place every visual immediately after the result paragraph that first interprets it.
    wide_burst(p[68]._p, [fig2, new_captions[0]])
    wide_burst(p[69]._p, [fig3, new_captions[1]])
    wide_burst(p[76]._p, [table_title, table, table_note])
    wide_burst(p[81]._p, [fig4, new_captions[2]])
    insert_after(p[85]._p, [fig5, new_captions[3]])

    # Extend the IEEE-numbered reference list without altering existing entries.
    cursor = p[107]._p
    for reference in REFERENCE_ADDITIONS:
        element = clone_with_text(p[107], reference)
        cursor.addnext(element)
        cursor = element

    # Force Word to balance the final two-column reference page.  This is a
    # separate unnumbered paragraph, so it cannot create a spurious [16].
    balance_break = deepcopy(p[39]._p)
    set_section_break_continuous(balance_break)
    body = document._body._element
    body.insert(len(body) - 1, balance_break)

    add_update_fields_setting(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)

    if embedded_hash(output, "word/media/image2.png") != FIGURE1_MEDIA_SHA256:
        raise RuntimeError("Figure 1 media changed unexpectedly")


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the clean IEEE SMC camera-ready manuscript candidate")
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--figures", type=Path, required=True)
    parser.add_argument("--table-csv", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build(args.source.resolve(), args.figures.resolve(), args.table_csv.resolve(), args.output.resolve())
    print(args.output)
    print(f"sha256={sha256(args.output.resolve())}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
